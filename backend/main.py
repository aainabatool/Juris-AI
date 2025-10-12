# # main.py
# from fastapi import FastAPI, Body, UploadFile, File
# from pydantic import BaseModel
# import os
# import getpass
# from langchain.chat_models import init_chat_model
# from langchain.prompts import PromptTemplate, FewShotPromptTemplate
# from langchain.chains import LLMChain
# from langchain.tools import Tool
# from langchain.agents import initialize_agent, AgentType
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# from docx import Document

# load_dotenv()

# # -------------------------------------------------
# # 🔹 API & Model Initialization
# # -------------------------------------------------
# app = FastAPI(
#     title="Pakistani Legal AI API",
#     description="Classifies and prioritizes Pakistani legal cases using LangChain & Gemini",
#     version="1.1"
# )

# if not os.environ.get("GOOGLE_API_KEY"):
#     os.environ["GOOGLE_API_KEY"] = getpass.getpass("Enter Google Gemini API Key: ")

# model = init_chat_model("gemini-2.5-flash", model_provider="google_genai")


# # -------------------------------------------------
# # 🔹 1. Legal Classification Chain
# # -------------------------------------------------
# classification_examples = [
#     {"case": "A person arrested under Section 302 PPC for murder. Bail application filed in Sessions Court.",
#      "classification": "Criminal Law"},
#     {"case": "میاں بیوی میں طلاق کا معاملہ۔ بیوی حق مہر اور کسٹڈی کا مطالبہ کر رہی ہے۔",
#      "classification": "Family Law"},
#     {"case": "FBR issued tax demand of Rs. 30 crore. Businessman filed appeal claiming excessive assessment.",
#      "classification": "Tax Law"},
#     {"case": "Writ petition under Article 199 challenging illegal detention by police. Fundamental rights violated.",
#      "classification": "Constitutional Law"},
#     {"case": "Government officer dismissed from service without inquiry. Filed appeal in Service Tribunal.",
#      "classification": "Service Law"},
#     {"case": "Landlord filed eviction suit against tenant for 10 months rent default. Property in Karachi.",
#      "classification": "Property Law"},
#     {"case": "50 factory workers terminated illegally. Labour Court ordered reinstatement. Employer appealed.",
#      "classification": "Labour Law"},
#     {"case": "Company director sued for breach of fiduciary duty and misappropriation of Rs. 50 million.",
#      "classification": "Corporate Law"}
# ]

# classification_example_template = """
# Case: {case}
# Classification: {classification}
# """

# classification_prompt = FewShotPromptTemplate(
#     examples=classification_examples,
#     example_prompt=PromptTemplate(
#         input_variables=["case", "classification"],
#         template=classification_example_template
#     ),
#     prefix="""You are an expert Pakistani legal case classifier with deep knowledge of Pakistan's legal system.
# Classify the case into one of the following categories:
# 1. Criminal Law
# 2. Civil Law
# 3. Family Law
# 4. Constitutional Law
# 5. Corporate Law
# 6. Labour Law
# 7. Property Law
# 8. Tax Law
# 9. Banking Law
# 10. Service Law
# 11. Islamic Law
# 12. Administrative Law
# 13. Environmental Law
# 14. Intellectual Property
# 15. Election Law

# Respond in this EXACT format:
# Category: [category name]
# Reasoning: [2-3 sentences explaining why this category fits]

# EXAMPLES:""",
#     suffix="""
# NOW CLASSIFY THIS CASE:
# {input}

# Category: 
# Reasoning:
# """,
#     input_variables=["input"]
# )

# classification_chain = LLMChain(llm=model, prompt=classification_prompt)


# # -------------------------------------------------
# # 🔹 2. Priority Assessment Chain
# # -------------------------------------------------
# priority_examples = [
#     {"case": "Murder accused arrested under Section 302 PPC. Bail hearing scheduled in 3 days. Accused in jail for 2 months.",
#      "priority": "High"},
#     {"case": "Property dispute over boundary wall between neighbors. Case pending for 5 years. No immediate harm.",
#      "priority": "Low"},
#     {"case": "Habeas corpus petition filed. Person detained illegally by police for 10 days without FIR. Family has no contact.",
#      "priority": "Critical"},
#     {"case": "Wife seeking maintenance for herself and 3 minor children. Husband not paying for 8 months. Children's education affected.",
#      "priority": "High"},
#     {"case": "Tax appeal against FBR assessment. Amount Rs. 5 lakh. Appeal deadline in 45 days.",
#      "priority": "Medium"},
# ]

# priority_example_template = """
# Case: {case}
# Priority: {priority}
# """

# priority_prompt = FewShotPromptTemplate(
#     examples=priority_examples,
#     example_prompt=PromptTemplate(
#         input_variables=["case", "priority"],
#         template=priority_example_template
#     ),
#     prefix="""You are an expert Pakistani legal case priority assessor.
# Decide the urgency level: CRITICAL, HIGH, MEDIUM, or LOW.

# Factors:
# - Life or liberty at risk → Critical
# - Fundamental rights violations → High or Critical
# - Deadlines within days/weeks → High
# - Routine matters → Low

# Respond in this EXACT format:
# Priority Level: [CRITICAL/HIGH/MEDIUM/LOW]
# Key Factors: [List of 3–4 points]
# Recommended Action Timeline: [e.g., Within 3 days, Within 2 weeks, etc.]
# Reasoning: [2–3 sentences]
# """,
#     suffix="""
# NOW ASSESS THIS CASE:
# {input}

# Priority Level: 
# Key Factors: 
# Recommended Action Timeline: 
# Reasoning:
# """,
#     input_variables=["input"]
# )

# priority_chain = LLMChain(llm=model, prompt=priority_prompt)


# # -------------------------------------------------
# # 🔹 3. Generic Legal Assistant Agent (Fallback)
# # -------------------------------------------------
# classification_tool = Tool(
#     name="Case Classifier",
#     func=lambda text: classification_chain.run(input=text),
#     description="Classifies a Pakistani legal case into the correct law category."
# )

# priority_tool = Tool(
#     name="Priority Assessor",
#     func=lambda text: priority_chain.run(input=text),
#     description="Assesses the urgency and importance of a legal case."
# )

# generic_prompt = """
# You are a friendly, highly intelligent Pakistani Legal AI assistant.
# If a case is ambiguous or does not match any predefined categories,
# you will engage the user conversationally, ask clarifying questions,
# and then decide which specialized tool (classifier or priority assessor)
# to use for final judgment.
# """

# agent = initialize_agent(
#     tools=[classification_tool, priority_tool],
#     llm=model,
#     agent_type=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
#     verbose=True
# )


# # -------------------------------------------------
# # 🔹 4. File Parsing Utility
# # -------------------------------------------------
# def extract_text_from_file(file: UploadFile) -> str:
#     ext = file.filename.lower().split(".")[-1]
#     if ext == "txt":
#         return file.file.read().decode("utf-8")

#     elif ext == "pdf":
#         reader = PdfReader(file.file)
#         text = ""
#         for page in reader.pages:
#             text += page.extract_text() or ""
#         return text.strip()

#     elif ext in ["docx", "doc"]:
#         doc = Document(file.file)
#         return "\n".join([p.text for p in doc.paragraphs])

#     else:
#         raise ValueError("Unsupported file type. Please upload .txt, .pdf, or .docx files.")


# # -------------------------------------------------
# # 🔹 5. Request Schemas
# # -------------------------------------------------
# class CaseInput(BaseModel):
#     text: str


# # -------------------------------------------------
# # 🔹 6. API Endpoints
# # -------------------------------------------------
# @app.post("/classify")
# def classify_case(case: CaseInput):
#     result = classification_chain.run(input=case.text)
#     return {"classification_result": result}


# @app.post("/priority")
# def assess_priority(case: CaseInput):
#     result = priority_chain.run(input=case.text)
#     return {"priority_result": result}


# @app.post("/analyze")
# def full_analysis(case: CaseInput):
#     classification = classification_chain.run(input=case.text)
#     priority = priority_chain.run(input=case.text)
#     return {"classification": classification, "priority": priority}


# @app.post("/agent")
# def legal_agent(case: CaseInput):
#     """Fallback AI agent that interacts conversationally if category is uncertain."""
#     response = agent.run(case.text)
#     return {"agent_response": response}


# @app.post("/upload")
# async def upload_file(file: UploadFile = File(...)):
#     """Extract text from uploaded legal file."""
#     try:
#         text = extract_text_from_file(file)
#         return {"filename": file.filename, "extracted_text": text[:2000]}  # preview first 2000 chars
#     except Exception as e:
#         return {"error": str(e)}


# # -------------------------------------------------
# # 🔹 Root
# # -------------------------------------------------
# @app.get("/")
# def root():
#     return {
#         "message": "Pakistani Legal AI API is running.",
#         "endpoints": ["/classify", "/priority", "/analyze", "/agent", "/upload"]
#     }


# main.py
from fastapi import FastAPI, Body, UploadFile, File
from pydantic import BaseModel
import os
import getpass
from langchain.chat_models import init_chat_model
from langchain.prompts import PromptTemplate, FewShotPromptTemplate
from langchain.chains import LLMChain
from langchain.tools import Tool
from langchain.agents import initialize_agent, AgentType
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
from fastapi.middleware.cors import CORSMiddleware

from typing import Optional, Dict, List, Any
from uuid import uuid4
import threading

# Import your agentic rag runner (assumes agentic_rag.py is in same folder)
from agentic_rag import run_agentic_rag

load_dotenv()

# -------------------------------------------------
# 🔹 API & Model Initialization
# -------------------------------------------------
app = FastAPI(
    title="Pakistani Legal AI API",
    description="Classifies and prioritizes Pakistani legal cases using LangChain & Gemini + Agentic RAG",
    version="1.2"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # or specify: ["http://localhost:5173"]
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if not os.environ.get("GOOGLE_API_KEY"):
    os.environ["GOOGLE_API_KEY"] = getpass.getpass("Enter Google Gemini API Key: ")

# Initialize the chat model (Gemini)
model = init_chat_model("gemini-2.5-flash", model_provider="google_genai")


# -------------------------------------------------
# 🔹 1. Legal Classification Chain
# -------------------------------------------------
classification_examples = [
    {"case": "A person arrested under Section 302 PPC for murder. Bail application filed in Sessions Court.",
     "classification": "Criminal Law"},
    {"case": "میاں بیوی میں طلاق کا معاملہ۔ بیوی حق مہر اور کسٹڈی کا مطالبہ کر رہی ہے۔",
     "classification": "Family Law"},
    {"case": "FBR issued tax demand of Rs. 30 crore. Businessman filed appeal claiming excessive assessment.",
     "classification": "Tax Law"},
    {"case": "Writ petition under Article 199 challenging illegal detention by police. Fundamental rights violated.",
     "classification": "Constitutional Law"},
    {"case": "Government officer dismissed from service without inquiry. Filed appeal in Service Tribunal.",
     "classification": "Service Law"},
    {"case": "Landlord filed eviction suit against tenant for 10 months rent default. Property in Karachi.",
     "classification": "Property Law"},
    {"case": "50 factory workers terminated illegally. Labour Court ordered reinstatement. Employer appealed.",
     "classification": "Labour Law"},
    {"case": "Company director sued for breach of fiduciary duty and misappropriation of Rs. 50 million.",
     "classification": "Corporate Law"}
]

classification_example_template = """
Case: {case}
Classification: {classification}
"""

classification_prompt = FewShotPromptTemplate(
    examples=classification_examples,
    example_prompt=PromptTemplate(
        input_variables=["case", "classification"],
        template=classification_example_template
    ),
    prefix="""You are an expert Pakistani legal case classifier with deep knowledge of Pakistan's legal system.
Classify the case into one of the following categories:
1. Criminal Law
2. Civil Law
3. Family Law
4. Constitutional Law
5. Corporate Law
6. Labour Law
7. Property Law
8. Tax Law
9. Banking Law
10. Service Law
11. Islamic Law
12. Administrative Law
13. Environmental Law
14. Intellectual Property
15. Election Law

Respond in this EXACT format:
Category: [category name]
Reasoning: [2-3 sentences explaining why this category fits]

EXAMPLES:""",
    suffix="""
NOW CLASSIFY THIS CASE:
{input}

Category: 
Reasoning:
""",
    input_variables=["input"]
)

classification_chain = LLMChain(llm=model, prompt=classification_prompt)


# -------------------------------------------------
# 🔹 2. Priority Assessment Chain
# -------------------------------------------------
priority_examples = [
    {"case": "Murder accused arrested under Section 302 PPC. Bail hearing scheduled in 3 days. Accused in jail for 2 months.",
     "priority": "High"},
    {"case": "Property dispute over boundary wall between neighbors. Case pending for 5 years. No immediate harm.",
     "priority": "Low"},
    {"case": "Habeas corpus petition filed. Person detained illegally by police for 10 days without FIR. Family has no contact.",
     "priority": "Critical"},
    {"case": "Wife seeking maintenance for herself and 3 minor children. Husband not paying for 8 months. Children's education affected.",
     "priority": "High"},
    {"case": "Tax appeal against FBR assessment. Amount Rs. 5 lakh. Appeal deadline in 45 days.",
     "priority": "Medium"},
]

priority_example_template = """
Case: {case}
Priority: {priority}
"""

priority_prompt = FewShotPromptTemplate(
    examples=priority_examples,
    example_prompt=PromptTemplate(
        input_variables=["case", "priority"],
        template=priority_example_template
    ),
    prefix="""You are an expert Pakistani legal case priority assessor.
Decide the urgency level: CRITICAL, HIGH, MEDIUM, or LOW.

Factors:
- Life or liberty at risk → Critical
- Fundamental rights violations → High or Critical
- Deadlines within days/weeks → High
- Routine matters → Low

Respond in this EXACT format:
Priority Level: [CRITICAL/HIGH/MEDIUM/LOW]
Key Factors: [List of 3–4 points]
Recommended Action Timeline: [e.g., Within 3 days, Within 2 weeks, etc.]
Reasoning: [2–3 sentences]
""",
    suffix="""
NOW ASSESS THIS CASE:
{input}

Priority Level: 
Key Factors: 
Recommended Action Timeline: 
Reasoning:
""",
    input_variables=["input"]
)

priority_chain = LLMChain(llm=model, prompt=priority_prompt)


# -------------------------------------------------
# 🔹 3. Tools: Classification, Priority, and RAG
# -------------------------------------------------
classification_tool = Tool(
    name="Case Classifier",
    func=lambda text: classification_chain.run(input=text),
    description="Classifies a Pakistani legal case into the correct law category."
)

priority_tool = Tool(
    name="Priority Assessor",
    func=lambda text: priority_chain.run(input=text),
    description="Assesses the urgency and importance of a legal case."
)

# Wrap agentic_rag runner as a Tool (calls your Chroma/RAG pipeline)
rag_tool = Tool(
    name="Legal RAG",
    func=lambda text: run_agentic_rag(text),
    description="Retrieval-augmented answer from legal knowledge core (Chroma DB). Use for deep legal Q&A."
)

# The conversational agent can choose which tool to use
agent = initialize_agent(
    tools=[classification_tool, priority_tool, rag_tool],
    llm=model,
    agent_type=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True
)


# -------------------------------------------------
# 🔹 4. File Parsing Utility
# -------------------------------------------------
def extract_text_from_file(file: UploadFile) -> str:
    ext = file.filename.lower().split(".")[-1]
    if ext == "txt":
        return file.file.read().decode("utf-8")

    elif ext == "pdf":
        reader = PdfReader(file.file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()

    elif ext in ["docx", "doc"]:
        doc = Document(file.file)
        return "\n".join([p.text for p in doc.paragraphs])

    else:
        raise ValueError("Unsupported file type. Please upload .txt, .pdf, or .docx files.")


# -------------------------------------------------
# 🔹 5. Session Memory (in-memory)
# -------------------------------------------------
# NOTE: This is an in-memory session store for simplicity.
# For production, replace with Redis / DB-backed session store.
_sessions_lock = threading.Lock()
_sessions: Dict[str, List[Dict[str, Any]]] = {}  # session_id -> list of messages (dict with role/text)


def _new_session_id() -> str:
    return str(uuid4())


def _ensure_session(session_id: Optional[str]) -> str:
    with _sessions_lock:
        if not session_id:
            sid = _new_session_id()
            _sessions[sid] = []
            return sid
        if session_id not in _sessions:
            _sessions[session_id] = []
        return session_id


def _append_message(session_id: str, role: str, text: str):
    with _sessions_lock:
        _sessions.setdefault(session_id, []).append({"role": role, "text": text})


def _get_history(session_id: str) -> List[Dict[str, str]]:
    return _sessions.get(session_id, [])


# -------------------------------------------------
# 🔹 6. Request Schemas
# -------------------------------------------------
class ChatInput(BaseModel):
    message: str
    session_id: Optional[str] = None


# -------------------------------------------------
# 🔹 7. Chat Endpoint (agent decides tool)
# -------------------------------------------------
@app.post("/chat")
def chat_endpoint(payload: ChatInput):
    """
    Chat endpoint that keeps session memory and lets the agent decide which tool to call.
    - If the agent calls "Legal RAG", it will call run_agentic_rag() (your RAG pipeline).
    - If the agent calls classifier/priority tools, those run as before.
    - Otherwise agent responds conversationally.
    Returns: { session_id, response, history }
    """
    session_id = _ensure_session(payload.session_id)
    user_msg = payload.message.strip()
    _append_message(session_id, "user", user_msg)

    # Build a single prompt text from session history for the agent.run
    # (agent.run expects a single text input; we give a concatenation of conversation)
    history = _get_history(session_id)
    # Build a conversational string; simple format:
    convo_text = "\n".join([f"{m['role'].upper()}: {m['text']}" for m in history])

    # Agent.run will internally choose tools (including "Legal RAG") as needed.
    try:
        agent_response = agent.run(convo_text)
    except Exception as e:
        # If agent failed, fallback to RAG directly
        rag_result = run_agentic_rag(user_msg)
        if isinstance(rag_result, dict):
            rag_result = rag_result.get("response", str(rag_result))

        rag_result = run_agentic_rag(user_msg)
        agent_response = rag_result["answer"] if isinstance(rag_result, dict) else str(rag_result)



    # Append assistant response to memory
    _append_message(session_id, "assistant", agent_response)

    return {"session_id": session_id, "response": agent_response, "history": _get_history(session_id)}


# -------------------------------------------------
# 🔹 8. Existing Endpoints (Classification / Priority / Upload)
# -------------------------------------------------
class CaseInput(BaseModel):
    text: str


@app.post("/classify")
def classify_case(case: CaseInput):
    result = classification_chain.run(input=case.text)
    return {"classification_result": result}


@app.post("/priority")
def assess_priority(case: CaseInput):
    result = priority_chain.run(input=case.text)
    return {"priority_result": result}


@app.post("/analyze")
def full_analysis(case: CaseInput):
    classification = classification_chain.run(input=case.text)
    priority = priority_chain.run(input=case.text)
    return {"classification": classification, "priority": priority}


@app.post("/agent")
def legal_agent(case: CaseInput):
    """Fallback AI agent that interacts conversationally if category is uncertain."""
    response = agent.run(case.text)
    return {"agent_response": response}


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Extract text from uploaded legal file."""
    try:
        text = extract_text_from_file(file)
        return {"filename": file.filename, "extracted_text": text[:2000]}  # preview first 2000 chars
    except Exception as e:
        return {"error": str(e)}


# -------------------------------------------------
# 🔹 Root
# -------------------------------------------------
@app.get("/")
def root():
    return {
        "message": "Pakistani Legal AI API is running.",
        "endpoints": ["/chat", "/classify", "/priority", "/analyze", "/agent", "/upload"]
    }
